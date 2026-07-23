"""Smart-fill user-uploaded IR DOCX templates (preserve letterhead/styles)."""

from __future__ import annotations

import io
import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from app.schemas import InvestigationReport
from app.services.ir_format import (
    allegation_export_line,
    conclusion_export_lines,
    facility_header_lines,
    sync_report_text,
)

SIZE_BODY = 12.0

CORE_SECTIONS = ("intake", "allegations", "process", "summary", "conclusion", "actions")
MIN_CORE_HEADINGS = 3

# First match wins; order matters for overlapping prefixes.
_SECTION_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("intake", re.compile(r"^intake\s+details\b", re.I)),
    # Header forms: Allegation(s): / Allegations: — not body "Allegation: Potential…"
    ("allegations", re.compile(r"^allegation(?:\(s\)|s)\s*:", re.I)),
    ("process", re.compile(r"^investigative\s+process\b", re.I)),
    ("summary", re.compile(r"^summary\s+of\s+findings\b", re.I)),
    ("conclusion", re.compile(r"^conclusion\b", re.I)),
    ("actions", re.compile(r"^actions\s*:?\s*$", re.I)),
]

_FACILITY_LABELS = [
    "Facility Address:",
    "Laboratory Director:",
    "CLIA Number:",
    "Credential Number:",
    "Medicare Number:",
    "Shell Number:",
    "Date(s) of Investigation:",
    "State Licensing Priority:",
    "Federal Certification Priority:",
]

_UNDERLINE_LABELS = {
    "pre-investigation activity",
    "investigation activity",
}
_BOLD_SUBHEADS = {
    "observations",
    "interviews",
    "document review",
}


@dataclass
class DetectedSection:
    key: str
    heading_index: int
    heading_text: str
    confidence: float = 1.0


@dataclass
class SectionMap:
    sections: list[DetectedSection]
    warnings: list[str]
    core_count: int

    def to_json(self) -> str:
        return json.dumps(
            {
                "sections": [asdict(s) for s in self.sections],
                "warnings": self.warnings,
                "core_count": self.core_count,
            }
        )

    @classmethod
    def from_json(cls, raw: str | None) -> SectionMap | None:
        if not raw:
            return None
        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            return None
        sections = [
            DetectedSection(
                key=s["key"],
                heading_index=int(s["heading_index"]),
                heading_text=s.get("heading_text") or "",
                confidence=float(s.get("confidence") or 1.0),
            )
            for s in (data.get("sections") or [])
        ]
        return cls(
            sections=sections,
            warnings=list(data.get("warnings") or []),
            core_count=int(data.get("core_count") or 0),
        )


class TemplateFillError(Exception):
    """Raised when a custom template cannot be filled safely."""

    def __init__(self, message: str, *, missing: list[str] | None = None):
        super().__init__(message)
        self.missing = missing or []


def _norm_heading(text: str) -> str:
    t = re.sub(r"\s+", " ", (text or "").strip())
    # Strip trailing instructional hint only (space before '('), keep Allegation(s):
    t = re.sub(r"\s+\(.*$", "", t).strip()
    return t


def _match_section_key(text: str) -> str | None:
    head = _norm_heading(text)
    if not head:
        return None
    # Do not treat body lines like "Allegation: Potential violation..." as section headers
    for key, pat in _SECTION_PATTERNS:
        if pat.search(head):
            return key
    return None


def detect_sections(doc: Document) -> SectionMap:
    """Scan body paragraphs for DOH-style section headings."""
    found: list[DetectedSection] = []
    seen: set[str] = set()
    for i, p in enumerate(doc.paragraphs):
        key = _match_section_key(p.text or "")
        if not key or key in seen:
            continue
        seen.add(key)
        found.append(
            DetectedSection(
                key=key,
                heading_index=i,
                heading_text=(p.text or "").strip()[:240],
                confidence=1.0,
            )
        )
    core = [s for s in found if s.key in CORE_SECTIONS]
    warnings: list[str] = []
    missing = [k for k in CORE_SECTIONS if k not in seen]
    if len(core) < MIN_CORE_HEADINGS:
        warnings.append(
            f"Found {len(core)} core section heading(s); need at least {MIN_CORE_HEADINGS}. "
            f"Missing: {', '.join(missing) if missing else 'none'}."
        )
    elif missing:
        warnings.append(f"Optional / unmatched headings: {', '.join(missing)}.")
    return SectionMap(sections=found, warnings=warnings, core_count=len(core))


def detect_sections_from_bytes(data: bytes) -> SectionMap:
    return detect_sections(Document(io.BytesIO(data)))


def detect_sections_from_path(path: Path | str) -> SectionMap:
    return detect_sections(Document(str(path)))


def _set_run_font(
    run,
    *,
    size_pt: float = SIZE_BODY,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool = False,
) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline:
        run.underline = WD_UNDERLINE.SINGLE
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _delete_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def _insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    """Insert a new paragraph immediately after `paragraph` and return it."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        _set_run_font(run)
    return new_para


def _clear_paragraph_text(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""
    if not paragraph.runs:
        paragraph.add_run("")


def _write_plain_after(anchor: Paragraph, lines: list[str]) -> Paragraph:
    cur = anchor
    for line in lines:
        cur = _insert_paragraph_after(cur, "")
        run = cur.add_run(line)
        _set_run_font(run)
    return cur


def _norm_label(line: str) -> str:
    return re.sub(r"\s+", " ", (line or "").strip().rstrip(":").lower())


def _write_process_after(anchor: Paragraph, steps: list[str]) -> Paragraph:
    cur = anchor
    for step in steps:
        text = (step or "").rstrip()
        key = _norm_label(text)
        cur = _insert_paragraph_after(cur, "")
        if key in _UNDERLINE_LABELS:
            body = text.strip()
            if body.endswith(":"):
                r1 = cur.add_run(body[:-1])
                _set_run_font(r1, bold=True, underline=True)
                r2 = cur.add_run(":")
                _set_run_font(r2, bold=True)
            else:
                r = cur.add_run(body)
                _set_run_font(r, bold=True, underline=True)
        elif key in _BOLD_SUBHEADS:
            r = cur.add_run(text)
            _set_run_font(r, bold=True)
        else:
            r = cur.add_run(text)
            _set_run_font(r)
    return cur


def _comparison_by_code(report: InvestigationReport) -> dict[str, Any]:
    out: dict[str, Any] = {}
    for c in report.comparisons or []:
        out[c.code] = c
        out[c.wac_id] = c
        # bare code without chapter prefix
        if "-" in c.code:
            out[c.code.split()[-1] if " " in c.code else c.code] = c
    return out


def allegation_lines_with_compare(report: InvestigationReport) -> list[str]:
    """Allegation export lines plus short Compare research notes."""
    by = _comparison_by_code(report)
    lines: list[str] = []
    for a in report.allegations:
        lines.append(allegation_export_line(a))
        comp = by.get(a.wac_code) or by.get(f"WAC {a.wac_code}")
        if not comp:
            continue
        notes: list[str] = []
        if comp.match_reason:
            notes.append(str(comp.match_reason).strip())
        subs = list(comp.matched_subsections or [])
        if subs:
            notes.append("Matched subsections: " + "; ".join(subs[:8]))
        if notes:
            lines.append("Evidence from comparison: " + " ".join(notes))
        lines.append("")  # spacer between allegations
    # drop trailing blank
    while lines and not lines[-1].strip():
        lines.pop()
    return lines


def _content_range(
    sections: list[DetectedSection], key: str, para_count: int
) -> tuple[int, int, int] | None:
    """Return (heading_index, content_start, content_end_exclusive) for key."""
    ordered = sorted(sections, key=lambda s: s.heading_index)
    idx = next((i for i, s in enumerate(ordered) if s.key == key), None)
    if idx is None:
        return None
    heading = ordered[idx].heading_index
    end = ordered[idx + 1].heading_index if idx + 1 < len(ordered) else para_count
    return heading, heading + 1, end


def _replace_section_body(
    doc: Document,
    heading_index: int,
    content_start: int,
    content_end: int,
    writer,
) -> None:
    """Delete paragraphs [content_start, content_end) and insert new body after heading."""
    # Delete from end so indices stay valid against live paragraph list
    for i in range(content_end - 1, content_start - 1, -1):
        if 0 <= i < len(doc.paragraphs):
            _delete_paragraph(doc.paragraphs[i])
    # Heading index may have shifted if we deleted before it — we only delete after heading
    if heading_index >= len(doc.paragraphs):
        return
    heading = doc.paragraphs[heading_index]
    writer(heading)


def _update_facility_fields(doc: Document, report: InvestigationReport) -> None:
    values = {label.lower(): value for label, value in facility_header_lines(report)}
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        for label in _FACILITY_LABELS:
            if text.lower().startswith(label.lower().rstrip(":").lower()) or text.lower().startswith(
                label.lower()
            ):
                # Normalize to "Label: value"
                value = values.get(label.lower(), "")
                _clear_paragraph_text(p)
                # Keep first run styling if possible
                run = p.runs[0] if p.runs else p.add_run()
                run.text = f"{label} {value}".rstrip()
                _set_run_font(run)
                break


def validate_section_map(section_map: SectionMap) -> None:
    if section_map.core_count < MIN_CORE_HEADINGS:
        missing = [k for k in CORE_SECTIONS if k not in {s.key for s in section_map.sections}]
        raise TemplateFillError(
            f"Template needs at least {MIN_CORE_HEADINGS} recognizable section headings "
            f"(Intake, Allegation(s), Investigative Process, Summary, Conclusion, Actions). "
            f"Found {section_map.core_count}. Missing: {', '.join(missing)}.",
            missing=missing,
        )


def smart_fill(template_path: Path | str, report: InvestigationReport) -> bytes:
    """Fill a user DOCX template with report content; preserve headers/footers."""
    report = sync_report_text(report)
    path = Path(template_path)
    if not path.is_file():
        raise TemplateFillError(f"Template file missing: {path}")

    doc = Document(str(path))
    section_map = detect_sections(doc)
    validate_section_map(section_map)

    _update_facility_fields(doc, report)

    # Build content writers per section
    intake = (report.intake_details or "").strip()
    process_steps = [str(s) for s in (report.investigative_process or [])]
    summary = (report.summary_of_findings or "").strip()
    conclusions = conclusion_export_lines(report)
    actions = (report.actions or "[To be determined after investigation]").strip()
    alleg_lines = allegation_lines_with_compare(report)

    writers = {
        "intake": lambda h: _write_plain_after(h, ["", intake] if intake else [""]),
        "allegations": lambda h: _write_plain_after(h, [""] + alleg_lines if alleg_lines else [""]),
        "process": lambda h: _write_process_after(
            _insert_paragraph_after(h, ""), process_steps
        ),
        "summary": lambda h: _write_plain_after(h, ["", summary] if summary else [""]),
        "conclusion": lambda h: _write_plain_after(
            h, [""] + [line for c in conclusions for line in (c, "")]
        ),
        "actions": lambda h: _write_plain_after(h, [actions] if actions else [""]),
    }

    # Replace from bottom to top so heading indices stay valid
    ordered = sorted(section_map.sections, key=lambda s: s.heading_index, reverse=True)
    para_count = len(doc.paragraphs)
    # Recompute ranges against original indices; apply bottom-up with fresh detect each time
    # Safer: re-detect after each replacement
    for sec in ordered:
        if sec.key not in writers:
            continue
        live = detect_sections(doc)
        rng = _content_range(live.sections, sec.key, len(doc.paragraphs))
        if not rng:
            continue
        heading_i, start, end = rng
        _replace_section_body(doc, heading_i, start, end, writers[sec.key])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_report_docx(
    report: InvestigationReport,
    *,
    template_path: Path | str | None = None,
    draft_label: str = "",
) -> bytes:
    """Export IR: custom smart-fill when template_path set, else built-in blank."""
    del draft_label
    if template_path:
        return smart_fill(template_path, report)
    from app.services.docx_export import build_investigation_docx

    return build_investigation_docx(report)
