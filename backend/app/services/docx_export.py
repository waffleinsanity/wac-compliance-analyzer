"""DOCX export for Investigation Reports - blank template styles + DOH run formatting."""

from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.schemas import InvestigationReport
from app.services.ir_blank import (
    ACTIONS_LABEL,
    ALLEGATION_HEADER,
    CONCLUSION_HEADER,
    INTAKE_LABEL,
    PROCESS_HEADER,
    SUMMARY_HEADER,
    TITLE,
    blank_docx_path,
    compose_actions_text,
    parse_actions_fields,
)
from app.services.ir_format import (
    allegation_export_line,
    conclusion_export_lines,
    facility_header_lines,
    sync_report_text,
)

STYLE_TITLE = "No Spacing"
STYLE_HEADER = "Header"
STYLE_BODY = "No Spacing"

SIZE_SECTION = 16.0
SIZE_BODY = 12.0
# Peer IRs indent allegation / intake body ~0.5in under section headings
BODY_INDENT = Inches(0.5)

_UNDERLINE_LABELS = {
    "pre-investigation activity",
    "investigation activity",
}
_BOLD_SUBHEADS = {
    "observations",
    "interviews",
    "document review",
}


def _clear_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _set_run_font(
    run,
    *,
    size_pt: float = SIZE_BODY,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool = False,
) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline:
        run.underline = WD_UNDERLINE.SINGLE
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _new_para(doc: Document, style: str):
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
    if p.runs:
        for run in p.runs:
            run.text = ""
    return p


def _add(
    doc: Document,
    text: str,
    style: str,
    *,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    size_pt: float = SIZE_BODY,
    center: bool = False,
    indent: bool = False,
) -> None:
    p = _new_para(doc, style)
    run = p.add_run(text or "")
    _set_run_font(run, size_pt=size_pt, bold=bold, italic=italic, underline=underline)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.left_indent = BODY_INDENT


def _add_blank(doc: Document, style: str = STYLE_BODY) -> None:
    _add(doc, "", style)


def _add_facility_line(doc: Document, label: str, value: str) -> None:
    """Bold label + regular value (blank Header style)."""
    p = _new_para(doc, STYLE_HEADER)
    r1 = p.add_run(f"{label} ")
    _set_run_font(r1, size_pt=SIZE_BODY, bold=True)
    r2 = p.add_run(value or "")
    _set_run_font(r2, size_pt=SIZE_BODY, bold=False)


def _parse_heading(label: str) -> tuple[str, str | None]:
    """Split 'Title: (hint…)' or 'Title (hint…)' into title + parenthetical hint."""
    text = (label or "").strip()
    m = re.match(r"^(.+?:)\s*(\(.*\))\s*$", text, flags=re.DOTALL)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    m = re.match(r"^(.+?)\s+(\(.*\))\s*$", text, flags=re.DOTALL)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return text, None


def _add_section_heading(doc: Document, label: str) -> None:
    """16pt bold title + 12pt italic hint — matches blank Intake / Process / Summary."""
    title, hint = _parse_heading(label)
    p = _new_para(doc, STYLE_BODY)
    title_text = f"{title} " if hint else title
    r = p.add_run(title_text)
    _set_run_font(r, size_pt=SIZE_SECTION, bold=True)
    if hint:
        r2 = p.add_run(hint)
        _set_run_font(r2, size_pt=SIZE_BODY, italic=True)


def _norm_label(line: str) -> str:
    return re.sub(r"\s+", " ", (line or "").strip().rstrip(":").lower())


def _add_process_line(doc: Document, line: str) -> None:
    text = (line or "").rstrip()
    key = _norm_label(text)
    if key in _UNDERLINE_LABELS:
        # Blank: bold+underline on words; colon bold only (not underlined)
        p = _new_para(doc, STYLE_BODY)
        body = text.strip()
        if body.endswith(":"):
            r1 = p.add_run(body[:-1])
            _set_run_font(r1, size_pt=SIZE_BODY, bold=True, underline=True)
            r2 = p.add_run(":")
            _set_run_font(r2, size_pt=SIZE_BODY, bold=True)
        else:
            r = p.add_run(body)
            _set_run_font(r, size_pt=SIZE_BODY, bold=True, underline=True)
        return
    if key in _BOLD_SUBHEADS:
        _add(doc, text, STYLE_BODY, bold=True, size_pt=SIZE_BODY)
        return
    from app.services.evidence_log import strip_trailing_superscripts

    body, marks = strip_trailing_superscripts(text)
    if marks:
        p = _new_para(doc, STYLE_BODY)
        p.paragraph_format.left_indent = BODY_INDENT
        run = p.add_run(body)
        _set_run_font(run, size_pt=SIZE_BODY)
        # Prefer Word superscript digits over Unicode marks for DOCX fidelity.
        digits = marks.translate(str.maketrans("⁰¹²³⁴⁵⁶⁷⁸⁹", "0123456789"))
        sup = p.add_run(digits)
        _set_run_font(sup, size_pt=SIZE_BODY)
        try:
            sup.font.superscript = True
        except Exception:
            pass
        return
    _add(doc, text, STYLE_BODY, size_pt=SIZE_BODY, indent=True)


def build_investigation_docx(
    report: InvestigationReport | dict[str, Any],
    *,
    draft_label: str = "Working draft — for investigator review",
    exhibits: list[Any] | None = None,
) -> bytes:
    """Build IR DOCX using blank template styles and DOH run formatting."""
    del draft_label
    if isinstance(report, dict):
        report = InvestigationReport.model_validate(report)
    report = sync_report_text(report)
    if exhibits:
        from app.services.evidence_log import annotate_process_with_exhibits

        report.investigative_process = annotate_process_with_exhibits(
            list(report.investigative_process or []),
            exhibits,
        )

    path = blank_docx_path()
    if not path.is_file():
        raise FileNotFoundError(f"Blank Investigation Report template missing: {path}")

    doc = Document(str(path))
    _clear_body(doc)

    # Title: 16pt bold underline, centered
    p = _new_para(doc, STYLE_TITLE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    _set_run_font(r, size_pt=SIZE_SECTION, bold=True, underline=True)
    _add_blank(doc, STYLE_BODY)

    # Investigation type (blank content-control dropdown under title)
    inv_type = (getattr(report, "subtitle", None) or "").strip()
    if inv_type:
        _add(doc, inv_type, STYLE_BODY, italic=True, center=True, size_pt=SIZE_BODY)
        _add_blank(doc, STYLE_BODY)

    for label, value in facility_header_lines(report):
        _add_facility_line(doc, label, value)

    _add_blank(doc, STYLE_BODY)
    _add_section_heading(doc, INTAKE_LABEL)
    _add_blank(doc, STYLE_BODY)
    _add(
        doc,
        (report.intake_details or "").strip(),
        STYLE_BODY,
        size_pt=SIZE_BODY,
        indent=True,
    )
    _add_blank(doc, STYLE_BODY)

    _add_section_heading(doc, ALLEGATION_HEADER)
    _add_blank(doc, STYLE_BODY)
    for i, a in enumerate(report.allegations, start=1):
        _add(
            doc,
            allegation_export_line(a, index=i),
            STYLE_BODY,
            size_pt=SIZE_BODY,
            indent=True,
        )
        _add_blank(doc, STYLE_BODY)

    _add_section_heading(doc, PROCESS_HEADER)
    _add_blank(doc, STYLE_BODY)
    for step in report.investigative_process or []:
        _add_process_line(doc, str(step))

    _add_blank(doc, STYLE_BODY)
    _add_section_heading(doc, SUMMARY_HEADER)
    _add_blank(doc, STYLE_BODY)
    from app.services.investigation import clean_summary_for_document

    _add(
        doc,
        clean_summary_for_document(report.summary_of_findings),
        STYLE_BODY,
        size_pt=SIZE_BODY,
        indent=True,
    )
    _add_blank(doc, STYLE_BODY)

    _add(doc, CONCLUSION_HEADER.strip(), STYLE_BODY, bold=True, size_pt=SIZE_SECTION)
    _add_blank(doc, STYLE_BODY)
    for line in conclusion_export_lines(report):
        _add(doc, line, STYLE_BODY, size_pt=SIZE_BODY, indent=True)
        _add_blank(doc, STYLE_BODY)

    _add(doc, ACTIONS_LABEL.strip(), STYLE_BODY, bold=True, size_pt=SIZE_SECTION)
    det, ref = parse_actions_fields(
        report.actions or "",
        determination=getattr(report, "action_determination", "") or "",
        referral=getattr(report, "action_referral", "") or "",
    )
    for line in compose_actions_text(det, ref).splitlines():
        _add(doc, line, STYLE_BODY, size_pt=SIZE_BODY, indent=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_deficiency_cite_sheet(report: InvestigationReport | dict[str, Any]) -> bytes:
    """Working cite sheet (internal). Prefer build_sod_docx for facility-facing SOD."""
    if isinstance(report, InvestigationReport):
        data = report.model_dump()
    else:
        data = report

    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Deficiency Cite Sheet (Working draft)")
    r.bold = True
    r.font.size = Pt(14)

    p = doc.add_paragraph(f"Case ID: {data.get('case_id') or '—'}")
    for run in p.runs:
        run.font.size = Pt(11)

    fi = data.get("facility_info") or {}
    doc.add_paragraph(f"Facility: {fi.get('facility_address') or '—'}")
    doc.add_paragraph(f"Credential: {fi.get('credential_number') or '—'}")
    doc.add_paragraph()

    sod = data.get("sod") or {}
    defs = sod.get("deficiencies") or []
    if defs:
        for d in defs:
            bp = doc.add_paragraph()
            br = bp.add_run(f"Cite: {d.get('regulation_cite') or ''}")
            br.bold = True
            doc.add_paragraph((d.get("based_on") or "")[:500])
            doc.add_paragraph()
    else:
        substantiated = [
            c
            for c in (data.get("conclusions") or [])
            if "substantiated" in (c.get("result") or "").lower()
        ]
        if not substantiated:
            doc.add_paragraph("No SOD deficiency blocks and no substantiated conclusions yet.")
        else:
            for c in substantiated:
                bp = doc.add_paragraph()
                br = bp.add_run(f"Cite: {c.get('wac_code') or ''}")
                br.bold = True
                doc.add_paragraph(c.get("allegation_text") or "")
                doc.add_paragraph(c.get("deficiency_details") or "Deficiency details pending.")
                doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_sod_docx(report: InvestigationReport | dict[str, Any] | None = None, *, exhibits: list[Any] | None = None) -> bytes:
    """Facility-facing SOD: fill Investigation SOD Template.docx in place.

    Preserves landscape page setup, DOH logo header, Word list bullets, bold titles,
    DRAFT watermark, and spacing from the blank. Identifier key is omitted.
    """
    from app.services.sod_template import build_sod_docx_bytes

    return build_sod_docx_bytes(report, exhibits=exhibits)


def build_sod_identifier_key(report: InvestigationReport | dict[str, Any]) -> bytes:
    """Internal-only identifier key - do not include in facility-facing packs by default."""
    if isinstance(report, InvestigationReport):
        data = report.model_dump()
    else:
        data = report
    sod = data.get("sod") or {}
    doc = Document()
    t = doc.add_paragraph()
    r = t.add_run("SOD Identifier Key (INTERNAL - do not send to facility)")
    r.bold = True
    r.font.size = Pt(14)
    for entry in sod.get("identifier_key") or []:
        doc.add_paragraph(
            f"{entry.get('kind') or 'Patient'} {entry.get('code') or ''}: "
            f"{entry.get('description') or ''}"
        )
    if not (sod.get("identifier_key") or []):
        doc.add_paragraph("No identifier mappings recorded yet.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
