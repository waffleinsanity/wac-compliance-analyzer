"""Fill Investigation SOD Template.docx from the Investigation Report.

The blank at data/templates/Investigation SOD Template.docx is the sole pack
shell (landscape, DOH logo header, Word lists, bold titles, spacing). Export
and preview return that file with case / deficiency fields filled in place.
Never invent statute text; regulation_text must come from report.sod (PDF-backed).
"""

from __future__ import annotations

import io
from copy import deepcopy
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

from app.config import settings
from app.schemas import InvestigationReport
from app.services.content_review import find_removal_spans
from app.services.sod_blank import (
    agency_display_line,
    blank_sod_docx_path,
    format_findings_column,
)

# Light yellow: auto-filled values investigators must verify (IR amber assist spirit).
_VERIFY_YELLOW = "FFFF99"
_DATE_PLACEHOLDER = "Click here to enter a date."

_LOGO_PNG_CANDIDATES = (
    settings.templates_dir / "sod-doh-logo.png",
    settings.project_root / "frontend" / "public" / "sod-doh-logo.png",
)


def _logo_png_bytes() -> bytes | None:
    for path in _LOGO_PNG_CANDIDATES:
        if path.is_file():
            return path.read_bytes()
    return None


def read_blank_sod_template_bytes() -> bytes:
    """Exact bytes of Investigation SOD Template.docx (unmodified)."""
    path = blank_sod_docx_path()
    if not path.is_file():
        raise FileNotFoundError(f"Investigation SOD Template missing: {path}")
    return path.read_bytes()


def _set_paragraph_text(
    paragraph,
    text: str,
    *,
    cites_by_no: dict[int, Any] | None = None,
) -> None:
    """Set paragraph text; trailing Unicode superscripts become Word superscript hyperlinks."""
    text = text if text is not None else ""
    from app.services.evidence_log import strip_trailing_superscripts

    body, marks = strip_trailing_superscripts(text)
    # Clear all existing runs
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = body
    else:
        paragraph.add_run(body)
    if not marks:
        return
    digits = marks.translate(str.maketrans("⁰¹²³⁴⁵⁶⁷⁸⁹", "0123456789"))
    while len(paragraph.runs) > 1:
        r = paragraph.runs[-1]
        r._r.getparent().remove(r._r)
    if cites_by_no:
        from app.services.evidence_cite import add_superscript_hyperlink

        for dig in digits:
            try:
                n = int(dig)
            except ValueError:
                continue
            cite = cites_by_no.get(n)
            if cite:
                add_superscript_hyperlink(
                    paragraph,
                    dig,
                    cite.pack_relpath,
                    tooltip=cite.tooltip,
                )
            else:
                sup = paragraph.add_run(dig)
                try:
                    sup.font.superscript = True
                except Exception:
                    pass
        return
    sup = paragraph.add_run(digits)
    try:
        sup.font.superscript = True
    except Exception:
        pass


def _annotate_deficiency_for_exhibits(deficiency: dict[str, Any] | Any, by_id: dict) -> dict[str, Any]:
    from app.services.evidence_log import annotate_finding_text_with_exhibits

    if hasattr(deficiency, "model_dump"):
        d = deficiency.model_dump()
    else:
        d = dict(deficiency or {})
    fins = []
    for f in d.get("findings") or []:
        fd = dict(f) if not isinstance(f, dict) else dict(f)
        fd["text"] = annotate_finding_text_with_exhibits(
            fd.get("text") or "",
            fd.get("evidence_ids") or [],
            by_id,
        )
        fins.append(fd)
    d["findings"] = fins
    items = []
    for it in d.get("items") or []:
        itd = dict(it) if not isinstance(it, dict) else dict(it)
        it_fins = []
        for f in itd.get("findings") or []:
            fd = dict(f) if not isinstance(f, dict) else dict(f)
            fd["text"] = annotate_finding_text_with_exhibits(
                fd.get("text") or "",
                fd.get("evidence_ids") or [],
                by_id,
            )
            it_fins.append(fd)
        itd["findings"] = it_fins
        items.append(itd)
    d["items"] = items
    return d


def _replace_in_paragraph(paragraph, old: str, new: str) -> bool:
    """Replace first occurrence of old across runs without moving text across SDTs."""
    if not old:
        return False
    segments: list[list] = []
    current: list = []
    for child in paragraph._p:
        tag = child.tag
        if tag == qn("w:r"):
            current.append(child)
        elif tag == qn("w:proofErr"):
            continue
        elif tag == qn("w:pPr"):
            continue
        else:
            if current:
                segments.append(current)
                current = []
    if current:
        segments.append(current)

    for seg in segments:
        t_nodes = []
        for r in seg:
            t_nodes.extend(r.iter(qn("w:t")))
        full = "".join(t.text or "" for t in t_nodes)
        if old not in full:
            continue
        updated = full.replace(old, new, 1)
        if not t_nodes:
            continue
        t_nodes[0].text = updated
        for t in t_nodes[1:]:
            t.text = ""
        return True
    return False


def _set_unique_cell_text(cell, text: str) -> None:
    text = text if text is not None else ""
    if cell.paragraphs:
        _set_paragraph_text(cell.paragraphs[0], text)
        for extra in cell.paragraphs[1:]:
            _set_paragraph_text(extra, "")
    else:
        cell.text = text


def _set_shd(element, fill: str = _VERIFY_YELLOW) -> None:
    """Set w:shd on a pPr or tcPr element (light yellow verify mark)."""
    for old in element.findall(qn("w:shd")):
        element.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def _paragraph_verify_shading(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    _set_shd(pPr)


def _cell_verify_shading(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    _set_shd(tcPr)


def _sdt_content_verify_shading(sdt) -> None:
    """Shade paragraphs inside an SDT content block (cover / meta dates)."""
    content = sdt.find(qn("w:sdtContent"))
    if content is None:
        return
    for p in content.iter(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p.insert(0, pPr)
        _set_shd(pPr)


def _findings_block_needs_verify(text: str) -> bool:
    """True when a findings block is assistive seed language (Based on / Failure to)."""
    body = (text or "").strip()
    if not body:
        return False
    if find_removal_spans(body):
        return True
    low = body.lower()
    return low.startswith("based on ") or low.startswith("failure to ")


def _set_findings_cell_with_verify(
    cell,
    deficiency: dict[str, Any] | Any,
    *,
    exhibit_by_id: dict | None = None,
    cites_by_no: dict[int, Any] | None = None,
) -> None:
    """Write findings as separate paragraphs; yellow-shade seed Based on / Failure to only."""
    d = deficiency
    if exhibit_by_id:
        d = _annotate_deficiency_for_exhibits(deficiency, exhibit_by_id)
    raw = format_findings_column(d)
    blocks = [b.strip() for b in raw.split("\n\n") if b.strip()]
    if not blocks:
        _set_unique_cell_text(cell, "")
        return
    while len(cell.paragraphs) < len(blocks):
        cell.add_paragraph()
    for i, block in enumerate(blocks):
        _set_paragraph_text(cell.paragraphs[i], block, cites_by_no=cites_by_no)
        if _findings_block_needs_verify(block):
            _paragraph_verify_shading(cell.paragraphs[i])
    for extra in cell.paragraphs[len(blocks) :]:
        _set_paragraph_text(extra, "")


def _unique_row_cells(row):
    cells = []
    for tc in row._tr.findall(qn("w:tc")):
        cells.append(_Cell(tc, row.table))
    return cells


def _replace_sdt_text(doc: Document, old: str, new: str, *, shade_verify: bool = False) -> int:
    """Replace placeholder text inside content controls (dates, etc.)."""
    count = 0
    for sdt in doc.element.body.iter(qn("w:sdt")):
        texts = list(sdt.iter(qn("w:t")))
        joined = "".join(t.text or "" for t in texts)
        if old not in joined:
            continue
        if texts:
            texts[0].text = joined.replace(old, new, 1)
            for t in texts[1:]:
                t.text = ""
            count += 1
            if shade_verify:
                _sdt_content_verify_shading(sdt)
    return count


def _report_dict(report: InvestigationReport | dict[str, Any] | None) -> dict[str, Any]:
    if report is None:
        return {}
    if isinstance(report, InvestigationReport):
        return report.model_dump()
    return dict(report)


def _ensure_sod_payload(data: dict[str, Any]) -> dict[str, Any]:
    """Build SOD from comparisons when report.sod is missing (preview/export path)."""
    sod = data.get("sod")
    if sod and (sod.get("deficiencies") or sod.get("facility_address") or sod.get("case_id")):
        return sod
    from app.services.sod_draft import build_sod_from_comparisons

    fi = data.get("facility_info") or {}
    built = build_sod_from_comparisons(
        data.get("comparisons") or [],
        case_id=data.get("case_id"),
        facility_address=fi.get("facility_address") or "",
        credential_number=fi.get("credential_number") or "",
        investigation_dates=fi.get("investigation_dates") or data.get("investigation_date") or "",
    )
    return built.model_dump() if built else {}


def _cover_name_and_address(facility_name: str, facility_address: str) -> tuple[str, str]:
    """Split Name / Address for the cover letter without duplicating a single line."""
    name = (facility_name or "").strip()
    addr_lines = [ln.strip() for ln in (facility_address or "").splitlines() if ln.strip()]
    if name:
        # Drop address lines that merely repeat the name.
        rest = [ln for ln in addr_lines if ln.lower() != name.lower()]
        return name, "\n".join(rest)
    if len(addr_lines) >= 2:
        return addr_lines[0], "\n".join(addr_lines[1:])
    if len(addr_lines) == 1:
        # Single IR facility line → Name only (do not copy into Address).
        return addr_lines[0], ""
    return "", ""


def _embed_png_logo(docx_bytes: bytes) -> bytes:
    """Swap WMF seal for PNG so browser preview (and Word) can show the logo."""
    png = _logo_png_bytes()
    if not png:
        return docx_bytes
    src = io.BytesIO(docx_bytes)
    out = io.BytesIO()
    with ZipFile(src, "r") as zin, ZipFile(out, "w", compression=ZIP_DEFLATED) as zout:
        for info in zin.infolist():
            name = info.filename
            data = zin.read(name)
            if name == "word/media/image1.wmf":
                continue
            if name == "[Content_Types].xml":
                text = data.decode("utf-8")
                if 'Extension="png"' not in text:
                    text = text.replace(
                        '<Default Extension="wmf" ContentType="image/x-wmf"/>',
                        '<Default Extension="wmf" ContentType="image/x-wmf"/>'
                        '<Default Extension="png" ContentType="image/png"/>',
                    )
                data = text.encode("utf-8")
            elif name == "word/_rels/document.xml.rels":
                text = data.decode("utf-8")
                text = text.replace('Target="media/image1.wmf"', 'Target="media/image1.png"')
                data = text.encode("utf-8")
            zout.writestr(info, data)
        zout.writestr("word/media/image1.png", png)
    return out.getvalue()


def fill_sod_template(
    report: InvestigationReport | dict[str, Any] | None = None,
    *,
    exhibits: list[Any] | None = None,
) -> Document:
    """Open Investigation SOD Template.docx and fill case / deficiency fields."""
    from app.services.evidence_cite import build_evidence_cites
    from app.services.evidence_log import exhibit_map_by_id
    from app.schemas import InvestigationReport as IRSchema

    data = _report_dict(report)
    sod = _ensure_sod_payload(data)
    fi = data.get("facility_info") or {}
    exhibit_by_id = exhibit_map_by_id(exhibits) if exhibits else {}
    cites_by_no: dict[int, Any] = {}
    if exhibits and report is not None:
        ir = report if isinstance(report, IRSchema) else IRSchema.model_validate(data)
        cites_by_no = {c.exhibit_no: c for c in build_evidence_cites(ir, exhibits)}

    facility_name = (sod.get("facility_name") or "").strip()
    facility_address = (sod.get("facility_address") or fi.get("facility_address") or "").strip()
    if facility_address.lower() in {"washington state", "n/a"}:
        # Same placeholder rule as IR facility highlight path.
        facility_address = "" if not facility_name else facility_address

    administrator = (
        (sod.get("administrator") or "").strip()
        or (fi.get("laboratory_director") or "").strip()
    )
    dates = (
        (sod.get("investigation_dates") or fi.get("investigation_dates") or data.get("investigation_date") or "")
        .strip()
    )
    investigator = (sod.get("investigator_number") or "").strip()
    poc_due = int(sod.get("poc_due_days") or 14)
    case_id = sod.get("case_id") or data.get("case_id") or ""
    license_no = sod.get("credential_number") or fi.get("credential_number") or ""
    services = (sod.get("agency_services_type") or "").strip()
    inspection = (sod.get("inspection_type") or "Investigation").strip() or "Investigation"

    path = blank_sod_docx_path()
    if not path.is_file():
        raise FileNotFoundError(f"Investigation SOD Template missing: {path}")

    doc = Document(str(path))
    paras = doc.paragraphs

    name_line, address_line = _cover_name_and_address(facility_name, facility_address)

    if name_line:
        _set_paragraph_text(paras[8], name_line)
        _paragraph_verify_shading(paras[8])
    if address_line:
        _set_paragraph_text(paras[9], address_line)
        _paragraph_verify_shading(paras[9])
    elif name_line:
        # Clear template "Address" placeholder so it does not echo Name.
        _set_paragraph_text(paras[9], "")

    if administrator:
        _set_paragraph_text(paras[11], f"Dear {administrator}:")
        _paragraph_verify_shading(paras[11])

    # Keep template wording; only swap placeholders.
    if name_line:
        _replace_in_paragraph(paras[13], "facility name", name_line)
    if dates:
        # Cover letter + meta date SDTs only. Never also write cell paragraph text
        # for the Investigation Start Date cell (that doubles the date string).
        _replace_sdt_text(doc, _DATE_PLACEHOLDER, dates, shade_verify=True)

    if poc_due != 14:
        _replace_in_paragraph(paras[15], "14 days", f"{poc_due} days")

    if investigator:
        _replace_in_paragraph(paras[30], "(surveyor number)", investigator)

    meta = doc.tables[0]
    agency_line = agency_display_line(facility_name or name_line, facility_address or name_line)
    r2 = _unique_row_cells(meta.rows[2])
    if len(r2) >= 1 and agency_line != "N/A":
        _set_unique_cell_text(r2[0], agency_line)
        _cell_verify_shading(r2[0])
    if len(r2) >= 3 and administrator:
        _set_unique_cell_text(r2[2], administrator)
        _cell_verify_shading(r2[2])

    r5 = _unique_row_cells(meta.rows[5])
    if len(r5) >= 1:
        _set_unique_cell_text(r5[0], inspection)
        _cell_verify_shading(r5[0])
    # Investigation Start Date: SDT fill above only (no _set_unique_cell_text).
    if len(r5) >= 3 and dates:
        _cell_verify_shading(r5[2])
    if len(r5) >= 5 and investigator:
        _set_unique_cell_text(r5[4], investigator)
        _cell_verify_shading(r5[4])

    r7 = _unique_row_cells(meta.rows[7])
    if len(r7) >= 1 and case_id:
        _set_unique_cell_text(r7[0], str(case_id))
        _cell_verify_shading(r7[0])
    if len(r7) >= 3 and license_no:
        _set_unique_cell_text(r7[2], license_no)
        _cell_verify_shading(r7[2])
    if len(r7) >= 5 and services:
        _set_unique_cell_text(r7[4], services)
        _cell_verify_shading(r7[4])

    def_table = doc.tables[1]
    deficiencies = sod.get("deficiencies") or []

    def _ensure_def_rows(needed: int) -> None:
        while len(def_table.rows) - 1 < needed:
            tbl = def_table._tbl
            last_tr = def_table.rows[-1]._tr
            tbl.append(deepcopy(last_tr))

    if deficiencies:
        _ensure_def_rows(len(deficiencies))
        for ri in range(1, len(def_table.rows)):
            for cell in def_table.rows[ri].cells:
                _set_unique_cell_text(cell, "")
        for i, d in enumerate(deficiencies):
            row = def_table.rows[i + 1].cells
            cite_parts = [
                (d.get("regulation_cite") or "").strip(),
                (d.get("regulation_text") or "").strip(),
            ]
            # Column 0 = PDF-backed statute authority: never yellow-shade.
            _set_unique_cell_text(row[0], "\n\n".join(p for p in cite_parts if p))
            _set_findings_cell_with_verify(
                row[1],
                d,
                exhibit_by_id=exhibit_by_id or None,
                cites_by_no=cites_by_no or None,
            )
            _set_unique_cell_text(row[2], "")

    return doc


def build_sod_docx_bytes(
    report: InvestigationReport | dict[str, Any] | None = None,
    *,
    exhibits: list[Any] | None = None,
) -> bytes:
    """Facility SOD: filled Investigation SOD Template (same bytes for preview + export)."""
    doc = fill_sod_template(report, exhibits=exhibits)
    buf = io.BytesIO()
    doc.save(buf)
    # Browser preview cannot paint WMF; PNG seal keeps preview/export visually aligned.
    return _embed_png_logo(buf.getvalue())
