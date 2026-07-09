"""Extract text from uploaded PDF/DOCX/TXT files."""

from __future__ import annotations

import io
from pathlib import Path

import fitz
from docx import Document


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return _pdf_bytes(data)
    if name.endswith(".docx"):
        return _docx_bytes(data)
    if name.endswith(".txt") or name.endswith(".md"):
        return data.decode("utf-8", errors="ignore")
    # Attempt docx then text
    try:
        return _docx_bytes(data)
    except Exception:
        return data.decode("utf-8", errors="ignore")


def extract_text_from_path(path: Path | str) -> str:
    path = Path(path)
    data = path.read_bytes()
    return extract_text_from_bytes(path.name, data)


def _pdf_bytes(data: bytes) -> str:
    doc = fitz.open(stream=data, filetype="pdf")
    parts = [doc[i].get_text() for i in range(len(doc))]
    return "\n".join(parts).strip()


def _docx_bytes(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    paras = [p.text for p in document.paragraphs if p.text.strip()]
    # Also tables
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n".join(paras).strip()
