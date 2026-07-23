"""Golden checks: exported IR matches blank DOCX styles and Pre-investigation shell."""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document

from app.schemas import (
    FacilityInfo,
    InvestigationAllegation,
    InvestigationConclusion,
    InvestigationReport,
)
from app.services.docx_export import build_investigation_docx
from app.services.ir_blank import BLANK_PROCESS_SKELETON, blank_docx_path
from app.services.ir_format import build_report_plain_text


def _sample_report() -> InvestigationReport:
    return InvestigationReport(
        title="Investigative Report",
        subtitle="State Investigation",
        investigation_date="07/23/2026",
        case_id="2025-11777",
        facility_info=FacilityInfo(
            facility_address="Washington State",
            credential_number="BHA.FS.61140707",
            investigation_dates="07/23/2026",
            state_licensing_priority="C",
        ),
        intake_details="The Department of Health (DOH) received a complaint alleging test.",
        allegation_preamble="jurisdiction",
        allegations=[
            InvestigationAllegation(
                case_category="BHA",
                wac_code="246-341-0410",
                wac_title="Administrator key responsibilities",
                allegation_text=(
                    "Potential violation of WAC 246-341-0410, Administrator key responsibilities, "
                    "by having failed to ensure day-to-day operation."
                ),
            ),
        ],
        investigative_process=list(BLANK_PROCESS_SKELETON),
        summary_of_findings="Pending investigation narrative.",
        conclusions=[
            InvestigationConclusion(
                wac_code="246-341-0410",
                allegation_text="x",
                result="Pending Investigation",
            ),
        ],
        actions="[To be determined after investigation]",
        comparisons=[],
        findings=[],
        report_text="",
        selected_count=1,
        duration_ms=1.0,
        document_preview="",
        regulatory_framework=[],
    )


def test_blank_template_has_expected_styles_and_preinvestigation():
    path = blank_docx_path()
    assert path.is_file()
    doc = Document(str(path))
    texts = [(p.style.name if p.style else "?", (p.text or "").strip()) for p in doc.paragraphs]
    joined = "\n".join(t for _, t in texts if t)
    assert "Investigative Report" in joined
    assert "Facility Address:" in joined
    assert "Pre-investigation Activity:" in joined
    assert "Investigation Activity:" in joined
    assert any(st == "Header" for st, t in texts if t.startswith("Facility Address"))
    assert any(st == "No Spacing" for st, t in texts if "Intake Details:" in t)


def test_plain_text_matches_blank_shell():
    text = build_report_plain_text(_sample_report())
    assert text.startswith("Investigative Report\n")
    assert "Facility Information" not in text
    assert "Facility Address: Washington State" in text
    assert "Laboratory Director:" in text
    assert "Intake Details: (List of concerns" in text
    assert "Allegation: Potential violation of WAC 246-341-0410" in text
    assert "Pre-investigation Activity:" in text
    assert "Investigation Activity:" in text
    assert "Allegation: The investigator found the facility pending determination of compliance" in text
    assert "Regulatory Framework" not in text
    assert "246-341-0410: Pending Investigation" not in text


def test_docx_export_uses_blank_styles_not_heading2_or_bullets():
    raw = build_investigation_docx(_sample_report())
    doc = Document(io.BytesIO(raw))
    styles = {p.style.name if p.style else "?" for p in doc.paragraphs}
    texts = [(p.style.name if p.style else "?", (p.text or "").strip()) for p in doc.paragraphs if (p.text or "").strip()]

    assert "Heading 2" not in styles
    assert "List Bullet" not in styles
    assert "Header" in styles
    assert "No Spacing" in styles

    joined = "\n".join(t for _, t in texts)
    assert "Facility Information" not in joined
    assert "State Investigation" not in joined
    assert "Working draft" not in joined
    assert "Regulatory Framework" not in joined
    assert any(st == "Header" and t.startswith("Facility Address:") for st, t in texts)
    assert any(t.startswith("Allegation: Potential violation") for _, t in texts)
    assert any("Pre-investigation Activity" in t for _, t in texts)
    assert any("found the facility pending determination of compliance" in t for _, t in texts)


def test_docx_export_doh_typography_hierarchy():
    """Section titles 16pt bold + italic hints; activity labels bold+underline."""
    raw = build_investigation_docx(_sample_report())
    doc = Document(io.BytesIO(raw))

    def find_para(substr: str):
        for p in doc.paragraphs:
            if substr in (p.text or ""):
                return p
        return None

    title = find_para("Investigative Report")
    assert title is not None and title.runs
    assert title.runs[0].bold is True
    assert title.runs[0].underline
    assert title.runs[0].font.size.pt == 16

    process = find_para("Investigative Process Included")
    assert process is not None and len(process.runs) >= 2
    assert process.runs[0].bold is True
    assert process.runs[0].font.size.pt == 16
    assert process.runs[1].italic is True
    assert process.runs[1].font.size.pt == 12

    pre = find_para("Pre-investigation Activity")
    assert pre is not None and pre.runs
    assert pre.runs[0].bold is True
    assert pre.runs[0].underline
    assert pre.runs[0].font.size.pt == 12

    obs = next(p for p in doc.paragraphs if (p.text or "").strip() == "Observations")
    assert obs.runs[0].bold is True
    assert not obs.runs[0].underline

    summary = find_para("Summary of Findings")
    assert summary is not None and len(summary.runs) >= 2
    assert summary.runs[0].bold is True
    assert summary.runs[0].font.size.pt == 16
    assert summary.runs[1].italic is True


def test_peer_examples_still_use_header_or_no_spacing(tmp_path: Path):
    """Sanity: at least one peer example keeps Header facility lines."""
    examples = Path(__file__).resolve().parents[2] / "data" / "examples"
    peers = list(examples.glob("*Peer*.docx")) + list(examples.glob("IR*.docx"))
    assert peers, "expected peer IR examples under data/examples"
    found_header = False
    for path in peers[:5]:
        doc = Document(str(path))
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            st = p.style.name if p.style else "?"
            if t.startswith("Facility Address") and st in {"Header", "No Spacing", "Normal"}:
                found_header = True
                break
        if found_header:
            break
    assert found_header
