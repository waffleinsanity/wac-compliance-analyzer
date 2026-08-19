"""Privacy scan / redact unit tests (complaint + evidence helpers)."""

from __future__ import annotations

from fastapi import HTTPException

from app.routers.cases import _scan_evidence_payload
from app.services.pii_gate import redact_text, scan_text


def test_scan_detects_ssn():
    text = "Patient SSN is 123-45-6789 in the intake note."
    scan = scan_text(text)
    assert scan["has_hits"] is True
    assert any(h["kind"] == "ssn" for h in scan["hits"])


def test_redact_clears_ssn():
    text = "Patient SSN is 123-45-6789 in the intake note."
    out = redact_text(text)
    assert "123-45-6789" not in out["redacted_text"]
    assert out["clean"] is True
    assert out["applied_count"] >= 1


def test_evidence_txt_auto_redacts_ssn():
    raw = b"Contact the complainant at 123-45-6789 for follow-up."
    data, note = _scan_evidence_payload("notes.txt", raw)
    assert b"123-45-6789" not in data
    assert "auto-redacted" in note


def test_evidence_clean_txt_passes():
    raw = b"Facility staffing on the evening shift was inadequate."
    data, note = _scan_evidence_payload("notes.txt", raw)
    assert data == raw
    assert "no Cat 3/4" in note


def test_evidence_docx_with_ssn_blocked():
    # Minimal bytes that extract_text may fail on — use a .txt-like path via md is covered;
    # for docx we simulate by checking HTTPException path with crafted extractable content.
    # Build a tiny DOCX via python-docx if available; otherwise skip-shaped assert on helper for pdf-like.
    from io import BytesIO
    from docx import Document

    buf = BytesIO()
    doc = Document()
    doc.add_paragraph("Resident MRN 998877 and SSN 321-54-9876 appear in this exhibit.")
    doc.save(buf)
    try:
        _scan_evidence_payload("exhibit.docx", buf.getvalue())
        raise AssertionError("expected HTTPException for PII docx")
    except HTTPException as exc:
        assert exc.status_code == 400
        assert "Category 3/4" in str(exc.detail)


def test_evidence_image_not_scanned():
    data, note = _scan_evidence_payload("photo.png", b"\x89PNG\r\n\x1a\nnot-real")
    assert data.startswith(b"\x89PNG")
    assert "image not text-scanned" in note


def test_wac_cite_is_not_a_phone():
    scan = scan_text("Potential violation of WAC 246-341-0410, Agency administration.")
    assert not any(h["kind"] == "phone" for h in scan["hits"])


def test_facility_policy_language_is_not_personal_name_or_org_phone():
    text = (
        "POLICIES AND PROCEDURES\n"
        "Infection Control Policy. Each individual receiving services must have an "
        "Individual Service Plan. The client may request Customer Service assistance. "
        "Contact the facility at (360) 555-0142. After hours crisis line: 1-800-273-8255.\n"
        "WAC 246-341-0425 Individual service record system."
    )
    scan = scan_text(text)
    kinds = {h["kind"] for h in scan["hits"]}
    assert "name" not in kinds
    assert "phone" not in kinds


def test_personal_phone_and_named_patient_still_detected():
    text = "Call the complainant at (360) 555-0199. Patient Jane Smith was admitted."
    scan = scan_text(text)
    kinds = {h["kind"] for h in scan["hits"]}
    assert "phone" in kinds
    assert "name" in kinds


def test_evidence_policy_docx_with_letterhead_is_not_blocked():
    from io import BytesIO
    from docx import Document

    buf = BytesIO()
    doc = Document()
    doc.add_paragraph(
        "Agency Policies and Procedures. Each individual must follow the Quality "
        "Management plan. Main office telephone (206) 555-0142. See WAC 246-341-0410."
    )
    doc.save(buf)
    data, note = _scan_evidence_payload("policies.docx", buf.getvalue())
    assert data
    assert "blocked" not in note.lower()


def test_evidence_docx_name_phone_only_is_not_blocked():
    from io import BytesIO
    from docx import Document

    buf = BytesIO()
    doc = Document()
    doc.add_paragraph("Approved by Dr. Jane Smith, Administrator. Fax (206) 555-0199.")
    doc.save(buf)
    data, note = _scan_evidence_payload("approval.docx", buf.getvalue())
    assert data
    assert "assistive" in note or "no Cat 3/4" in note
