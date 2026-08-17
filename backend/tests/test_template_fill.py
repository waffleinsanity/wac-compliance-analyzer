"""Smart-fill custom IR DOCX templates."""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document

from app.schemas import (
    FacilityInfo,
    InvestigationAllegation,
    InvestigationConclusion,
    InvestigationReport,
    WACComparison,
)
from app.services.ir_blank import BLANK_PROCESS_SKELETON, blank_docx_path
from app.services.template_fill import (
    TemplateFillError,
    detect_sections,
    detect_sections_from_bytes,
    smart_fill,
)


def _sample_report(**kwargs) -> InvestigationReport:
    base = dict(
        title="Investigative Report",
        subtitle="On-site State Investigation",
        investigation_date="07/23/2026",
        case_id="2025-11777",
        facility_info=FacilityInfo(
            facility_address="123 Test Ave",
            credential_number="BHA.FS.61140707",
            investigation_dates="07/23/2026",
            state_licensing_priority="C",
        ),
        intake_details="CUSTOM_INTAKE_MARKER: complaint narrative for fill test.",
        allegation_preamble="jurisdiction",
        allegations=[
            InvestigationAllegation(
                case_category="BHA",
                wac_code="246-341-0410",
                wac_title="Administrator key responsibilities",
                allegation_text=(
                    "Potential violation of WAC 246-341-0410, Administrator key responsibilities."
                ),
            ),
        ],
        investigative_process=list(BLANK_PROCESS_SKELETON),
        summary_of_findings="CUSTOM_SUMMARY_MARKER: pending narrative.",
        conclusions=[
            InvestigationConclusion(
                wac_code="246-341-0410",
                allegation_text="x",
                result="Pending Investigation",
            ),
        ],
        actions="[To be determined after investigation]",
        comparisons=[
            WACComparison(
                wac_id="WAC 246-341-0410",
                code="246-341-0410",
                title="Administrator key responsibilities",
                chapter="246-341",
                hierarchy_path="246-341-0410",
                wac_text="text",
                wac_summary="summary",
                allegation_draft="draft",
                matched_subsections=["246-341-0410(1)"],
                match_reason="COMPARE_REASON_MARKER: facility operations oversight.",
                match_score=0.9,
            ),
        ],
        findings=[],
        report_text="",
        selected_count=1,
        duration_ms=1.0,
        document_preview="",
        regulatory_framework=[],
    )
    base.update(kwargs)
    return InvestigationReport(**base)


def _add_header_letterhead(doc: Document, text: str) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def test_detect_sections_on_blank_template():
    path = blank_docx_path()
    assert path.is_file()
    sm = detect_sections(Document(str(path)))
    keys = {s.key for s in sm.sections}
    assert "intake" in keys
    assert "process" in keys
    assert sm.core_count >= 3


def test_smart_fill_preserves_letterhead_and_updates_body(tmp_path: Path):
    src = blank_docx_path()
    doc = Document(str(src))
    marker = "LETTERHEAD_UNIQUE_MARKER_XYZ"
    _add_header_letterhead(doc, marker)
    fixture = tmp_path / "custom_shell.docx"
    doc.save(str(fixture))

    raw = smart_fill(fixture, _sample_report())
    out = Document(io.BytesIO(raw))
    header_text = "\n".join(p.text for p in out.sections[0].header.paragraphs)
    assert marker in header_text

    body = "\n".join(p.text for p in out.paragraphs)
    assert "CUSTOM_INTAKE_MARKER" in body
    assert "CUSTOM_SUMMARY_MARKER" in body
    assert "COMPARE_REASON_MARKER" in body
    assert "Pre-investigation Activity" in body
    assert "123 Test Ave" in body


def test_missing_headings_raises(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("Cover only")
    path = tmp_path / "bad.docx"
    doc.save(str(path))

    sm = detect_sections_from_bytes(path.read_bytes())
    assert sm.core_count < 3

    with pytest.raises(TemplateFillError) as ei:
        smart_fill(path, _sample_report())
    msg = str(ei.value).lower()
    assert "at least 3" in msg or "missing" in msg
